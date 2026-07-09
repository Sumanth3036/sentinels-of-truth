import re
from datetime import datetime
from pathlib import Path

from docx import Document

from app.config import OUTPUT_DIR


def _slugify(title: str) -> str:
    slug = re.sub(r"[^\w\s-]", "", title.lower())
    slug = re.sub(r"[\s_-]+", "-", slug).strip("-")
    return slug[:60] or "document"


def generate_docx(title: str, content: dict) -> str:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=0)

    for section in content.get("sections", []):
        heading = section.get("heading", "Section")
        doc.add_heading(heading, level=1)

        if "bullets" in section and section["bullets"]:
            for item in section["bullets"]:
                doc.add_paragraph(str(item), style="List Bullet")
        elif "body" in section and section["body"]:
            doc.add_paragraph(str(section["body"]))

    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"{_slugify(title)}-{timestamp}.docx"
    filepath = OUTPUT_DIR / filename
    doc.save(str(filepath))

    return str(filepath).replace("\\", "/")
